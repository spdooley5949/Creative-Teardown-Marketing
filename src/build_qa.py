"""Q&A prep sheet for the Friday presentation. Two pages, printable."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor as RGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import sys

INK  = RGB(0x1F, 0x29, 0x33)
MID  = RGB(0x44, 0x51, 0x5F)
GREY = RGB(0x69, 0x74, 0x82)
SIG  = RGB(0xC2, 0x41, 0x0C)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.7)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10)
st.paragraph_format.space_after = Pt(0)


def para(text="", *, size=10, bold=False, italic=False, color=INK,
         before=0, after=0, font="Calibri", align=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size, r.font.bold, r.font.italic, r.font.name = Pt(size), bold, italic, font
    r.font.color.rgb = color
    return p


def rich(parts, *, before=0, after=0, indent=0, size=10):
    """parts = [(text, bold, color), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for t, b, c in parts:
        r = p.add_run(t)
        r.font.size, r.font.bold, r.font.name = Pt(size), b, "Calibri"
        r.font.color.rgb = c
    return p


# ---- header ---------------------------------------------------------------
para("CREATIVE TEARDOWN — QUESTION PREP", size=9, bold=True, color=SIG, after=2)
para("Shane Dooley  ·  Friday presentation  ·  Answers you already have",
     size=9, color=GREY, after=8)

para("Lead with the one-liner. The backup is only if they push. "
     "Every number below is in the workbook and traceable to specific ads.",
     size=9, italic=True, color=MID, after=12)


def qa(n, question, lead, backup, trap=None):
    rich([(f"{n}.  ", True, SIG), (question, True, INK)], before=9, after=3, size=10.5)
    rich([("Lead with:  ", True, SIG), (lead, False, INK)], indent=0.25, after=3)
    rich([("If pushed:  ", True, GREY), (backup, False, MID)], indent=0.25, after=3)
    if trap:
        rich([("Don't:  ", True, GREY), (trap, True, GREY)], indent=0.25, after=2)


qa(1,
   "How do you know the tags are right? An AI did the classification.",
   "A teammate audited all 101 ads row by row against the copy and found 42 errors. "
   "We applied every one, including two inside our own headline count.",
   "Tags are frozen to a committed file, so the numbers reproduce exactly between runs and "
   "corrections are permanent. Both matrices reconcile against the underlying tagged data "
   "with zero discrepancies, and the Tagged Ads sheet shows every ad with its tags, so any "
   "number traces back to the specific ads that produced it. The audit is in the repo as "
   "docs/tagging-audit.md.",
   "claim the tagging is perfect. Say it was audited, and that the audit found real errors.")

qa(2,
   "Search ads aren't where brands position themselves. Isn't this the wrong data?",
   "Largely correct, and we measured it: only 16 of 101 ads are brand-level. The rest is "
   "product feed and discounts. That's why we classified every ad by type.",
   "This is a teardown of search messaging, not of brand strategy as a whole. Brand positioning "
   "lives in video and social, which the Transparency Center won't give us as text. The tool "
   "runs with a brand-only flag that rebuilds the analysis on brand-level copy alone. The "
   "convergence finding holds in both cuts.",
   "get defensive. The honest scoping is the strongest version of this answer.")

qa(3,
   "You have a 1,979-ad dataset sitting in your own repo. Why analyse 101?",
   "We chose verified over large. The 101 are hand-transcribed, human-audited, and "
   "independently replicated. The 1,979 are none of those things yet.",
   "Merging the big file would have cost us the replication, because that argument depends on "
   "two datasets staying separate. Only 81 of its 1,979 rows carry an audited tag. It's merged "
   "and credited, and it's the obvious next step, but not the basis for a claim we make on "
   "Friday.",
   "apologise for the sample size. This was a decision, not a limitation.")

qa(4,
   "If community is such an opportunity, why hasn't anyone taken it?",
   "One brand has. Gymshark carries belonging in half its ads through coached couch-to-5K "
   "and half-marathon programmes. Ten brands have ceded the lane to it.",
   "And even Gymshark doesn't prove it. No participation number, no completion rate, no member "
   "count. The claim exists; the evidence doesn't. That's the gap: not an unclaimed territory, "
   "an unproven one, and proof is what a competitor can't copy by reformulating a fabric.",
   "say 'nobody does community.' That's the version our own audit disproved.")

qa(5,
   "Patagonia has three ads. Isn't your sample too thin to say anything?",
   "Patagonia genuinely runs three search text ads. That's the finding, not the limitation.",
   "Matrix cells are percentages of each brand's own ads, with the denominator shown, so a "
   "three-ad brand and an eighteen-ad brand stay comparable. A small sample does carry a wide "
   "error bar, which is why our headline rests on the category pattern across 101 ads rather "
   "than on any single brand's row.",
   "let them equate 'few ads collected' with 'few ads exist.' Those are different claims.")

qa(6,
   "Claude did the work. What did you actually do?",
   "I decided what to measure, caught what the model got wrong, and killed five findings "
   "before they reached this room.",
   "A false proof gap, a false sustainability whitespace, a keyword trend that was Adidas "
   "loyalty boilerplate, a five-year trend line that was really a change in sample composition, "
   "and 42 tagging errors a teammate surfaced. The model produced a confident wrong answer at "
   "every one of those points. The judgment about what counts as community, what counts as "
   "proof, and which dataset to trust was not automatable.",
   "be defensive or oversell the tool. The corrections are the credential.")

qa(7,
   "Were these ads actually running on 12 August?",
   "They're every text ad the Transparency Center listed for those advertisers that day. "
   "We didn't filter by date or open each ad to verify run status.",
   "That's why we say 'listed,' not 'live.' It doesn't change the finding, because theme "
   "distribution across the ads a brand has run is still a valid read of its messaging. The "
   "distinction is stated in the README and the methodology.",
   None)

qa(8,
   "How would you know if you were wrong?",
   "We tried to find out. We deliberately added Patagonia, Tracksmith and Fabletics to "
   "break the finding, and ran it against a second dataset we didn't collect.",
   "Patagonia did kill an earlier claim that sustainability was open ground. Community survived: "
   "across four brands appearing in both datasets, collected by different people using different "
   "methods, the community figures agree within 0 to 4 points. That's the test, and we ran it "
   "before presenting rather than after being challenged.",
   None)

# ---- numbers block --------------------------------------------------------
doc.add_page_break()
para("NUMBERS, IF YOU NEED THEM", size=9, bold=True, color=SIG, after=6)

rows = [
    ("The finding", "Performance 59% of ads · Comfort 58% · Versatility 51% · Style 36%"),
    ("", "Price/Value 34% · Innovation 14% · Sustainability 7% · Community 6%"),
    ("Community detail", "6 ads total. Gymshark 5 (50% of its ads), Alo Yoga 1. Ten brands at zero."),
    ("Proof", "33% of ads carry hard evidence. 67% carry none."),
    ("", "Alo 21,615 reviews · Vuori 13,593 · Nike 6,660 · UA 305 · Arc'teryx 24 · Tracksmith 22"),
    ("Discounting", "Under Armour runs price language in 86% of its ads. Fabletics 100%."),
    ("Ad mix", "16 of 101 brand-level. 62 product feed, 23 promotional."),
    ("Sample", "101 ads, 12 brands, listed 12 Aug 2026. lululemon 18 down to Patagonia 3."),
    ("Replication", "1,786-ad second dataset, 8 brands, 2021-2026, collected by a teammate."),
    ("", "4 brands overlap. Mean gap 8 points. Community agrees within 0-4 points."),
    ("Audit", "All 101 rows reviewed by Jessica Smith. 42 corrections applied."),
    ("Also in repo", "1,979-ad master list, 16 brands, deduplicated. Not used for these numbers."),
]
for label, val in rows:
    rich([((label + "  ") if label else "", True, INK), (val, False, MID)],
         indent=0.05, after=2, size=9.5)

para("THE THREE THINGS TO SAY IF EVERYTHING ELSE FAILS", size=9, bold=True,
     color=SIG, before=14, after=5)
for t in [
    "Every brand is selling the same three things, and the language is interchangeable.",
    "Two in three ads make a claim and prove nothing.",
    "One brand owns belonging, nobody contests it, and even the owner doesn't prove it.",
]:
    rich([("•  ", False, SIG), (t, True, INK)], indent=0.1, after=3, size=10)

para("Repo: github.com/spdooley5949/Creative-Teardown-Marketing",
     size=8.5, color=GREY, before=14)

doc.save(sys.argv[1])
print("wrote", sys.argv[1])
